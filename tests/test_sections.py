"""Dynamic PRE/MAIN/POST sections: a kind may repeat, in any order."""
import re


def _make_testcase(client, story_code="EX-700"):
    create = client.post("/stories", data={"display_code": story_code, "title": "Payments"}, follow_redirects=False)
    story_id = create.headers["location"].rstrip("/").split("/")[-1]
    client.post(f"/stories/{story_id}/phases", data={"type": "SIT"})
    page = client.get(f"/stories/{story_id}")
    phase_id = page.text.split('/subtasks/new')[0].split('/phases/')[-1]
    sub = client.post(
        f"/phases/{phase_id}/subtasks",
        data={"display_code": "S-1", "title": "Exec", "subtask_type": "EXECUTION"},
        follow_redirects=False,
    )
    subtask_id = sub.headers["location"].rstrip("/").split("/")[-1]
    client.post(f"/subtasks/{subtask_id}/testcases", data={"display_code": "TC-1", "title": "Flow"})
    detail = client.get(f"/subtasks/{subtask_id}")
    return re.search(r"/testcases/(\d+)/execute", detail.text).group(1)


def _section_ids(client, testcase_id):
    page = client.get(f"/testcases/{testcase_id}/execute")
    return re.findall(r"/sections/(\d+)/delete", page.text)


def _section_labels(client, testcase_id):
    """Section names in rendered order. Only section cards carry the index
    span, so the Description panel is excluded automatically."""
    page = client.get(f"/testcases/{testcase_id}/execute")
    return re.findall(r"<span data-section-index>\d+</span>\. ([^<]+)</div>", page.text)


def test_new_testcase_starts_with_the_three_default_sections(client):
    testcase_id = _make_testcase(client)
    assert _section_labels(client, testcase_id) == ["Pre Condition", "Main Test", "Post Condition"]


def test_sections_can_repeat_and_keep_the_order_they_were_added(client):
    testcase_id = _make_testcase(client, "EX-701")
    client.post(f"/testcases/{testcase_id}/sections", data={"kind": "MAIN"})
    client.post(f"/testcases/{testcase_id}/sections", data={"kind": "POSTCONDITION"})

    assert _section_labels(client, testcase_id) == [
        "Pre Condition", "Main Test", "Post Condition", "Main Test", "Post Condition",
    ]


def test_steps_belong_to_their_own_section_instance(client):
    testcase_id = _make_testcase(client, "EX-702")
    client.post(f"/testcases/{testcase_id}/sections", data={"kind": "MAIN"})
    first_main, second_main = _section_ids(client, testcase_id)[1], _section_ids(client, testcase_id)[3]

    client.post(f"/testcases/{testcase_id}/sections/{first_main}/steps", data={"step_text": "first main step"})
    client.post(f"/testcases/{testcase_id}/sections/{second_main}/steps", data={"step_text": "second main step"})

    page = client.get(f"/testcases/{testcase_id}/execute").text
    # Each step numbers from 1 within its own section.
    assert page.index("first main step") < page.index("second main step")
    assert page.count('name="step_text"') == 2


def test_deleting_a_section_removes_only_its_own_steps(client):
    testcase_id = _make_testcase(client, "EX-703")
    client.post(f"/testcases/{testcase_id}/sections", data={"kind": "MAIN"})
    ids = _section_ids(client, testcase_id)
    first_main, second_main = ids[1], ids[3]
    client.post(f"/testcases/{testcase_id}/sections/{first_main}/steps", data={"step_text": "keep me"})
    client.post(f"/testcases/{testcase_id}/sections/{second_main}/steps", data={"step_text": "remove me"})

    client.post(f"/testcases/{testcase_id}/sections/{second_main}/delete")

    page = client.get(f"/testcases/{testcase_id}/execute").text
    assert "keep me" in page
    assert "remove me" not in page
    assert len(_section_ids(client, testcase_id)) == 3


def test_export_renders_every_section_in_order(client, tmp_path):
    from docx import Document

    testcase_id = _make_testcase(client, "EX-704")
    client.post(f"/testcases/{testcase_id}/sections", data={"kind": "MAIN"})
    ids = _section_ids(client, testcase_id)
    client.post(f"/testcases/{testcase_id}/sections/{ids[1]}/steps", data={"step_text": "main one"})
    client.post(f"/testcases/{testcase_id}/sections/{ids[3]}/steps", data={"step_text": "main two"})

    response = client.get(f"/testcases/{testcase_id}/export-docx")
    assert response.status_code == 200
    out = tmp_path / "out.docx"
    out.write_bytes(response.content)

    doc = Document(str(out))
    headings = [p.text.strip() for p in doc.paragraphs if p.style.name == "Heading 1" and p.text.strip()]
    assert headings == ["PRE CONDITION", "MAIN TEST", "POST CONDITION", "MAIN TEST"]

    step_texts = [t.cell(0, 5).text for t in doc.tables[1:] if t.cell(0, 5).text]
    assert step_texts == ["main one", "main two"]


def test_image_zip_disambiguates_repeated_section_kinds(client):
    import io
    import zipfile
    import base64

    png = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk"
        "+A8AAQUBAScY42YAAAAASUVORK5CYII="
    )
    testcase_id = _make_testcase(client, "EX-705")
    client.post(f"/testcases/{testcase_id}/sections", data={"kind": "MAIN"})
    ids = _section_ids(client, testcase_id)
    for section_id, text in ((ids[1], "Check A"), (ids[3], "Check B")):
        client.post(f"/testcases/{testcase_id}/sections/{section_id}/steps", data={"step_text": text})
    page = client.get(f"/testcases/{testcase_id}/execute").text
    for step_id in re.findall(r"/steps/(\d+)/edit", page):
        client.post(
            f"/testcases/{testcase_id}/steps/{step_id}/screenshot",
            files={"file": ("p.png", png, "image/png")},
        )

    names = zipfile.ZipFile(io.BytesIO(client.get(f"/testcases/{testcase_id}/export-images").content)).namelist()
    # Each block's own heading letter keeps repeated kinds apart.
    assert names == ["B.MAIN-TEST_1.Check A.png", "D.MAIN-TEST_1.Check B.png"]


def test_sections_can_be_reordered(client):
    testcase_id = _make_testcase(client, "EX-706")
    pre, main, post = _section_ids(client, testcase_id)

    # Drag Post Condition to the front.
    response = client.post(
        f"/testcases/{testcase_id}/sections/reorder",
        data={"order": f"{post},{pre},{main}"},
        follow_redirects=False,
    )
    assert response.status_code == 303
    assert _section_labels(client, testcase_id) == ["Post Condition", "Pre Condition", "Main Test"]
    assert _section_ids(client, testcase_id) == [post, pre, main]


def test_reorder_rejects_ids_from_another_testcase(client):
    first = _make_testcase(client, "EX-707")
    second = _make_testcase(client, "EX-708")
    intruder = _section_ids(client, second)[0]
    original = _section_ids(client, first)

    response = client.post(
        f"/testcases/{first}/sections/reorder",
        data={"order": ",".join(original[:2] + [intruder])},
    )
    assert response.status_code == 422
    # Nothing was applied.
    assert _section_ids(client, first) == original


def test_reorder_survives_in_the_export(client, tmp_path):
    from docx import Document

    testcase_id = _make_testcase(client, "EX-709")
    pre, main, post = _section_ids(client, testcase_id)
    client.post(f"/testcases/{testcase_id}/sections/reorder", data={"order": f"{main},{post},{pre}"})

    response = client.get(f"/testcases/{testcase_id}/export-docx")
    out = tmp_path / "reordered.docx"
    out.write_bytes(response.content)
    doc = Document(str(out))
    headings = [p.text.strip() for p in doc.paragraphs if p.style.name == "Heading 1" and p.text.strip()]
    assert headings == ["MAIN TEST", "POST CONDITION", "PRE CONDITION"]
