from docx import Document

from app.docx.builder import build_docx


class _Step:
    def __init__(self, step_no, section, text, expected, actual, screenshots=None):
        self.step_no = step_no
        self.section = section
        self.step_text = text
        self.expected_result = expected
        self.actual_result = actual
        self.screenshots = screenshots or []


class _Screenshot:
    def __init__(self, file_path):
        self.file_path = file_path


class _Enum:
    def __init__(self, value):
        self.value = value

    @property
    def label(self):
        return self.value.replace("_", " ")


class _Section:
    def __init__(self, kind, steps):
        self.kind = kind
        self.steps = steps


def _sections_from(steps):
    """Group flat steps into sections, preserving first-seen order."""
    order, grouped = [], {}
    for step in steps:
        if step.section not in grouped:
            grouped[step.section] = []
            order.append(step.section)
        grouped[step.section].append(step)
    return [_Section(kind, sorted(grouped[kind], key=lambda s: s.step_no)) for kind in order]


class _Phase:
    def __init__(self, story, type_value):
        self.story = story
        self.type = _Enum(type_value)


class _Subtask:
    def __init__(self, story, title, phase_type):
        self.title = title
        self.phase = _Phase(story, phase_type)


class _Story:
    def __init__(self, title, display_code="EX-142"):
        self.title = title
        self.display_code = display_code


class _TestCase:
    def __init__(self, subtask, steps):
        self.subtask = subtask
        self.steps = steps
        self.display_code = "TC-1"
        self.title = "Verify top-up RO balance"
        self.tester = "Andri Firman Nurvianto"
        self.test_date = "2026-08-26"
        self.test_priority = "High"
        self.test_type = "Functional"
        self.channel = "Mobile App"
        self.iteration = "1"
        self.balance_before = "Rp. -"
        self.balance_after = "Rp. -"
        self.usage = "Rp. -"
        self.remark = ""
        self.data_test = "msisdn: 62812"
        self.status = _Enum("PASS")

    # The builder reads `sections`; tests still assign a flat `steps` list.
    @property
    def steps(self):
        return [step for section in self.sections for step in section.steps]

    @steps.setter
    def steps(self, value):
        self.sections = _sections_from(value)


def _make_testcase(steps):
    from app.models import StepSection

    story = _Story("Payments")
    subtask = _Subtask(story, "SIT Login Flow", "SIT")
    return _TestCase(subtask, steps), StepSection


def test_build_docx_header_and_single_step(tmp_path):
    tc, StepSection = _make_testcase([])
    tc.steps = [
        _Step(1, StepSection.PRECONDITION, "pre text", "pre expected", "pre actual"),
        _Step(1, StepSection.MAIN, "main text", "main expected", "main actual"),
        _Step(1, StepSection.POSTCONDITION, "post text", "post expected", "post actual"),
    ]
    output_path = str(tmp_path / "out.docx")
    build_docx(tc, output_path)

    doc = Document(output_path)
    # Use ground-truth row-scoped access (immune to flat-index bugs) for key header fields
    header_table = doc.tables[0]
    # Project identifies the task (story), Scenario the test case itself.
    assert header_table.rows[0].cells[3].text == "EX-142 - Payments"  # project (row 0)
    assert header_table.rows[1].cells[3].text == "TC-1 - Verify top-up RO balance"  # scenario (row 1)
    assert header_table.rows[4].cells[3].text == "SIT"  # environment (row 4, affected by flat-index bug if using .cell())
    assert header_table.rows[8].cells[3].text == "1"  # iteration (row 8, affected by flat-index bug)
    assert header_table.rows[13].cells[3].text == ""  # remark (row 13, affected by flat-index bug)
    assert header_table.rows[14].cells[-1].text == "msisdn: 62812"  # data_test (row 14, uses fallback to last cell)

    # Step blocks
    assert doc.tables[1].cell(0, 5).text == "pre text"
    assert doc.tables[2].cell(1, 2).text == "main actual"
    assert doc.tables[3].cell(1, 5).text == "post expected"


def test_build_docx_clones_tables_for_multiple_steps(tmp_path):
    tc, StepSection = _make_testcase([])
    tc.steps = [
        _Step(1, StepSection.MAIN, "step one", "e1", "a1"),
        _Step(2, StepSection.MAIN, "step two", "e2", "a2"),
        _Step(3, StepSection.MAIN, "step three", "e3", "a3"),
    ]
    output_path = str(tmp_path / "out2.docx")
    build_docx(tc, output_path)

    doc = Document(output_path)
    # Sections are explicit now: these steps define a single MAIN section, so
    # the export carries the header table plus one block per step and nothing
    # for kinds the test case does not have.
    assert len(doc.tables) == 4  # header + 3 MAIN blocks
    main_texts = [t.cell(0, 5).text for t in doc.tables if len(t.rows) == 4 and t.cell(0, 5).text.startswith("step")]
    assert main_texts == ["step one", "step two", "step three"]


def test_build_docx_inserts_screenshots(tmp_path):
    import base64

    png_path = tmp_path / "shot.png"
    png_path.write_bytes(base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk"
        "+A8AAQUBAScY42YAAAAASUVORK5CYII="
    ))
    tc, StepSection = _make_testcase([])
    tc.steps = [_Step(1, StepSection.MAIN, "step", "e", "a", screenshots=[_Screenshot(str(png_path))])]
    output_path = str(tmp_path / "out3.docx")
    build_docx(tc, output_path)

    doc = Document(output_path)
    assert len(doc.inline_shapes) == 1


def test_build_docx_screenshots_do_not_accumulate_across_cloned_step_tables(tmp_path):
    """Regression test for cloned step-block image accumulation.

    3 MAIN steps; screenshots on steps 1 and 2 only (step 3 has none). Each
    step's cloned table must contain exactly its own screenshots -- not the
    screenshots of steps that came before it. A bug where clones are made
    from an already-filled table (instead of a pristine snapshot) causes
    images to bleed forward into later steps' tables.
    """
    import base64

    png_path = tmp_path / "shot.png"
    png_path.write_bytes(base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk"
        "+A8AAQUBAScY42YAAAAASUVORK5CYII="
    ))
    tc, StepSection = _make_testcase([])
    tc.steps = [
        _Step(1, StepSection.MAIN, "step one", "e1", "a1", screenshots=[_Screenshot(str(png_path))]),
        _Step(2, StepSection.MAIN, "step two", "e2", "a2", screenshots=[_Screenshot(str(png_path))]),
        _Step(3, StepSection.MAIN, "step three", "e3", "a3"),
    ]
    output_path = str(tmp_path / "out_accum.docx")
    build_docx(tc, output_path)

    doc = Document(output_path)
    # Isolate the 3 MAIN step-block tables in document order (4-row step
    # blocks whose step-text cell starts with "step").
    step_tables = [
        t for t in doc.tables
        if len(t.rows) == 4 and t.cell(0, 5).text.startswith("step")
    ]
    assert [t.cell(0, 5).text for t in step_tables] == ["step one", "step two", "step three"]

    def _drawing_count(table):
        # Count w:drawing elements within just this table's own XML subtree,
        # so images are attributed to the correct table rather than counted
        # globally (a global doc.inline_shapes count would not distinguish
        # which step a bled-forward image landed in).
        return len(table._tbl.xpath(".//*[local-name()='drawing']"))

    counts = [_drawing_count(t) for t in step_tables]
    assert counts == [1, 1, 0]


def test_build_docx_preserves_data_test_field_despite_cell_index_anomaly(tmp_path):
    """Regression test: data_test field must be written even when row 14 col 3 is inaccessible.

    The template has a structural anomaly: row 14 column 3 raises IndexError when accessed via
    table.cell(), requiring a fallback to column 2. This test verifies the value is not silently
    dropped but actually written to the document.
    """
    tc, StepSection = _make_testcase([])
    tc.steps = [_Step(1, StepSection.MAIN, "step", "e", "a")]
    tc.data_test = "msisdn: 62812"
    output_path = str(tmp_path / "out_data_test.docx")
    build_docx(tc, output_path)

    doc = Document(output_path)
    # Verify the data_test value appears somewhere in the header table (exact column not guaranteed)
    header_table = doc.tables[0]
    header_text_found = False
    for row in header_table.rows:
        for cell in row.cells:
            if "msisdn: 62812" in cell.text.lower():
                header_text_found = True
                break
        if header_text_found:
            break
    assert header_text_found, "data_test field value 'msisdn: 62812' not found in header table"


def test_build_docx_handles_unembeddable_screenshot_without_crashing(tmp_path):
    """Regression test: a screenshot file add_picture can't embed must not abort the export.

    Screenshot format is deliberately never validated at upload time, so a referenced
    "screenshot" can be anything -- e.g. a non-image file, or an image format python-docx
    can't embed (such as WebP). build_docx must complete successfully and leave a
    traceable placeholder in that step's table instead of raising.
    """
    bad_path = tmp_path / "not_an_image.webp"
    bad_path.write_bytes(b"not actually an image")

    tc, StepSection = _make_testcase([])
    tc.steps = [_Step(1, StepSection.MAIN, "step", "e", "a", screenshots=[_Screenshot(str(bad_path))])]
    output_path = str(tmp_path / "out_bad_screenshot.docx")

    # Must not raise.
    build_docx(tc, output_path)

    doc = Document(output_path)
    step_table = next(t for t in doc.tables if len(t.rows) == 4 and t.cell(0, 5).text == "step")
    assert "not_an_image.webp" in step_table.cell(3, 0).text
    assert len(step_table._tbl.xpath(".//*[local-name()='drawing']")) == 0


def _png_bytes():
    import struct
    import zlib

    def chunk(tag, data):
        return struct.pack(">I", len(data)) + tag + data + struct.pack(">I", zlib.crc32(tag + data))

    w = h = 8
    raw = (b"\x00" + bytes((90, 120, 200)) * w) * h
    return (
        b"\x89PNG\r\n\x1a\n"
        + chunk(b"IHDR", struct.pack(">IIBBBBB", w, h, 8, 2, 0, 0, 0))
        + chunk(b"IDAT", zlib.compress(raw, 9))
        + chunk(b"IEND", b"")
    )


def test_build_docx_unwraps_content_controls_and_writes_test_date_once(tmp_path):
    """The template ships Word content controls (a date picker on Test Date,
    dropdowns elsewhere). The date control wraps the value CELL itself, so
    leaving it in place both hid the real cell from python-docx and left its
    stale placeholder in the row -- the date rendered twice. Export must
    contain no content controls and exactly one Test Date value.
    """
    import zipfile

    import lxml.etree as ET

    tc, StepSection = _make_testcase([])
    tc.steps = [_Step(1, StepSection.MAIN, "step", "e", "a")]
    tc.test_date = "2026-08-26"
    output_path = str(tmp_path / "out_controls.docx")
    build_docx(tc, output_path)

    ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    root = ET.fromstring(zipfile.ZipFile(output_path).read("word/document.xml"))
    assert root.findall(".//" + ns + "sdt") == [], "content controls must be unwrapped"

    header = Document(output_path).tables[0]
    date_row = header.rows[3]
    assert len(date_row.cells) == 4, "Test Date row must expose its real value cell"
    # Stored ISO date is rendered long-form for the artifact.
    assert date_row.cells[3].text == "Wednesday, 26 August 2026"
    # The stale placeholder from the date picker must be gone, and the value
    # must not also appear in the narrow spacer cell.
    row_text = [c.text for c in date_row.cells]
    assert row_text.count("Wednesday, 26 August 2026") == 1
    # The date picker's own placeholder must not survive alongside it.
    assert not any("11 August 2026" in t for t in row_text)


def test_build_docx_renders_data_test_as_bullet_list(tmp_path):
    tc, StepSection = _make_testcase([])
    tc.steps = [_Step(1, StepSection.MAIN, "step", "e", "a")]
    tc.data_test = "msisdn: 62812\nSID: 8117369\namount: 50000"
    output_path = str(tmp_path / "out_bullets.docx")
    build_docx(tc, output_path)

    ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    cell = Document(output_path).tables[0].rows[14].cells[3]
    paragraphs = [p for p in cell.paragraphs if p.text.strip()]
    assert [p.text for p in paragraphs] == ["msisdn: 62812", "SID: 8117369", "amount: 50000"]
    for paragraph in paragraphs:
        p_pr = paragraph._p.find(ns + "pPr")
        assert p_pr is not None and p_pr.find(ns + "numPr") is not None, (
            "each Data Test line must carry real Word bullet numbering"
        )


def test_build_docx_screenshots_are_18cm_wide_and_centered(tmp_path):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Emu

    png_path = tmp_path / "shot.png"
    png_path.write_bytes(_png_bytes())

    tc, StepSection = _make_testcase([])
    tc.steps = [_Step(1, StepSection.MAIN, "step", "e", "a", screenshots=[_Screenshot(str(png_path))])]
    output_path = str(tmp_path / "out_shot_size.docx")
    build_docx(tc, output_path)

    doc = Document(output_path)
    assert len(doc.inline_shapes) == 1
    assert round(Emu(doc.inline_shapes[0].width).cm, 2) == 18.0

    step_table = next(t for t in doc.tables if len(t.rows) == 4 and t.cell(0, 5).text == "step")
    shot_cell = step_table.cell(3, 0)
    assert shot_cell.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER


def test_project_and_scenario_codes_are_tracker_hyperlinks(tmp_path):
    """The code half of Project/Scenario links back to the ticket."""
    import zipfile

    import lxml.etree as ET

    tc, StepSection = _make_testcase([])
    tc.steps = [_Step(1, StepSection.MAIN, "step", "e", "a")]
    output_path = str(tmp_path / "out_links.docx")
    build_docx(tc, output_path)

    ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    rel_ns = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"
    archive = zipfile.ZipFile(output_path)
    root = ET.fromstring(archive.read("word/document.xml"))
    rels = {r.get("Id"): r.get("Target") for r in ET.fromstring(archive.read("word/_rels/document.xml.rels"))}

    links = {
        "".join(t.text or "" for t in link.iter(ns + "t")): rels.get(link.get(rel_ns + "id"))
        for link in root.findall(".//" + ns + "hyperlink")
    }
    assert links == {
        "EX-142": "https://collabs.xlsmart.co.id/browse/EX-142",
        "TC-1": "https://collabs.xlsmart.co.id/browse/TC-1",
    }

    # The visible text is still "<code> - <title>".
    header = Document(output_path).tables[0]
    assert header.rows[0].cells[3].text == "EX-142 - Payments"
    assert header.rows[1].cells[3].text == "TC-1 - Verify top-up RO balance"
